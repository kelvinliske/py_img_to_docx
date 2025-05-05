import os 
import pytesseract
 #usado em windows - colocar o caminho onde foi instalado - em outro OS comentar a linha de baixo
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def limpar_palavra_cortada(linha):
    # Remove palavras com hífen no fim da linha
    if linha.strip().endswith('-'):
        return ''
    return linha

#definição de formatação na extração do texto
def extrair_texto_para_word(imagem_path, destino, fonte_nome="Times New Roman", tamanho_fonte=10, tamanho_cabecalho=13):
    imagem = Image.open(imagem_path)
    texto = pytesseract.image_to_string(imagem, lang='por')

    doc = Document()
    nome_imagem = os.path.splitext(os.path.basename(imagem_path))[0]

    linhas = texto.strip().split("\n")
    
    for i, linha in enumerate(linhas):
        linha_limpa = limpar_palavra_cortada(linha.strip())
        if not linha_limpa:
            continue

        paragrafo = doc.add_paragraph()
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragrafo.add_run(linha_limpa)

        # Cabeçalho: primeira linha
        run.font.size = Pt(tamanho_cabecalho if i == 0 else tamanho_fonte)
        run.font.name = fonte_nome
        paragrafo.paragraph_format.line_spacing = 1.0

    caminho_arquivo = os.path.join(destino, f"{nome_imagem}.docx")
    doc.save(caminho_arquivo)
    print(f"Documento criado: {caminho_arquivo}")

def processar_pasta_imagens(caminho_pasta, destino, extensoes_validas=[".png", ".jpg", ".jpeg"]):
    if not os.path.exists(destino):
        os.makedirs(destino)

    for arquivo in os.listdir(caminho_pasta):
        if any(arquivo.lower().endswith(ext) for ext in extensoes_validas):
            imagem_path = os.path.join(caminho_pasta, arquivo)
            extrair_texto_para_word(imagem_path, destino)

# Caminhos
caminho_das_imagens = r"C:\\Users\\kliske\\Documents\\Kelvin\\Files\\Project files"
pasta_de_saida = r"C:\\Users\\kliske\\Documents\\Kelvin\\Files\\Project files\\Docs"

processar_pasta_imagens(caminho_das_imagens, pasta_de_saida)
